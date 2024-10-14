import streamlit as st
import json
import random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import zipfile
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

# Add a selection box for choosing the JSON file
json_option = st.selectbox(
    "Chọn đề cương",
    ("Đề cương lớp 10", "Đề cương lớp 11")
)

# Map the selection to the corresponding JSON file and headings
json_file_map = {
    "Đề cương lớp 10": {
        "file": "output_2.json",
        "headings": {
            "part1": "PHẦN I. TRẮC NGHIỆM ",
            "part2": "PHẦN II. TỰ LUẬN "
        }
    },
    "Đề cương lớp 11": {
        "file": "output.json",
        "headings": {
            "part1": "PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. Thí sinh trả lời câu hỏi từ câu 1 đến câu 24. Mỗi câu hỏi thí sinh chỉ chọn một phương án.",
            "part2": "PHẦN II. Câu trắc nghiệm đúng, sai: Thí sinh trả lời từ câu 1 đến câu 4, trong mỗi ý a, b, c, d ở mỗi câu thí sinh chọn đúng hoặc sai."
        }
    }
}

selected_json_file = json_file_map[json_option]["file"]
selected_headings = json_file_map[json_option]["headings"]

# Read the selected JSON file
with open(selected_json_file, 'r', encoding='utf-8') as file:
    data = json.load(file)

st.title("Tạo đề thi")

# Create widgets to select the number of questions for each section
selections = {}
for bai, content in data.items():
    st.header(bai)
    for phan, questions in content.items():
        key = f"{bai}_{phan}"
        num_questions = len(questions)
        selections[key] = st.number_input(f"Số câu hỏi cho {bai} {phan}", min_value=0, max_value=num_questions, value=0)

# Add widget to select the number of exams
num_exams = st.number_input("Số lượng đề cần tạo", min_value=1, max_value=10, value=1)


def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def format_question(doc, question_text, is_multiple_choice=True):
    # Split question and answers
    parts = question_text.split('\n')
    question = parts[0]
    answers = parts[1:]

    # Add question
    p = doc.add_paragraph(style='Normal')
    p.add_run(question)
    p.paragraph_format.space_after = Pt(0)  # Remove space after question

    if is_multiple_choice:
        # Create table for multiple choice answers
        if answers:
            table = doc.add_table(rows=len(answers), cols=1)
            table.allow_autofit = False
            table.width = Inches(6.5)

            for i, answer in enumerate(answers):
                cell = table.cell(i, 0)
                cell.text = answer.strip()
                cell_para = cell.paragraphs[0]
                cell_para.paragraph_format.space_after = Pt(0)
                for run in cell_para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                set_cell_border(
                    cell,
                    top={"sz": 0, "val": "none"},
                    bottom={"sz": 0, "val": "none"},
                    start={"sz": 0, "val": "none"},
                    end={"sz": 0, "val": "none"},
                )
    else:
        # For essay questions, add answers as normal paragraphs
        for answer in answers:
            p = doc.add_paragraph(answer.strip(), style='Normal')
            p.paragraph_format.space_after = Pt(0)

        # Add 10 lines for answers using tab stops
        for _ in range(10):
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.add_run('\t')  # Add a tab to create the dotted line
            p.paragraph_format.space_after = Pt(12)


def count_pages(doc):
    return len(doc.sections)


def add_horizontal_line(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'thick')
    bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def select_questions(data, selections):
    exam_questions_part1 = []
    exam_questions_part2 = []
    for key, num_selected in selections.items():
        bai, phan = key.split('_')
        questions = data[bai][phan]
        selected_questions = random.sample(list(questions.items()), num_selected)
        if phan == "Phần 1":
            exam_questions_part1.extend([(bai, phan, q_num, q_text) for q_num, q_text in selected_questions])
        elif phan == "Phần 2":
            exam_questions_part2.extend([(bai, phan, q_num, q_text) for q_num, q_text in selected_questions])

    random.shuffle(exam_questions_part1)
    random.shuffle(exam_questions_part2)
    return exam_questions_part1, exam_questions_part2

def create_custom_style(doc, name, font_name, font_size, bold=False):
    style = doc.styles.add_style(name, 1)
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    return style

if st.button("Tạo đề thi"):
    # Create a buffer to save the ZIP file
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for exam_number in range(1, num_exams + 1):
            # Select questions for this exam
            exam_questions_part1, exam_questions_part2 = select_questions(data, selections)

            # Create Word file
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            heading_style = create_custom_style(doc, 'CustomHeading', 'Times New Roman', 14, bold=True)
            # Adjust page margins
            section = doc.sections[0]
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

            # Add header table
            header_table = doc.add_table(rows=1, cols=2)
            header_table.allow_autofit = False
            header_table.width = Inches(8)  # Adjust as needed

            # Left cell
            left_cell = header_table.cell(0, 0)
            left_cell.width = Inches(4)
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align vertically
            left_para.add_run("SỞ GD&ĐT HÀ NỘI\n").bold = True
            left_para.add_run("TRƯỜNG THPT PHÚC LỢI\n").bold = True
            left_para.add_run("---------------\n")

            # Right cell
            right_cell = header_table.cell(0, 1)
            right_cell.width = Inches(4)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            right_para.add_run(f"KIỂM TRA GIỮA KÌ 1 - KHỐI {10 if '10' in json_option else 11}\n").bold = True
            right_para.add_run(f"NĂM HỌC 2024 - 2025\n").bold = True
            right_para.add_run("MÔN: LỊCH SỬ\n").bold = True
            right_para.add_run("Thời gian làm bài: 45 phút\n")
            right_para.add_run("(không kể thời gian phát đề)")

            # Remove border from table
            for row in header_table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top={"sz": 0, "val": "none"},
                        bottom={"sz": 0, "val": "none"},
                        start={"sz": 0, "val": "none"},
                        end={"sz": 0, "val": "none"},
                    )

            # Add name and code fields
            fields = doc.add_paragraph()
            fields.add_run("Họ và tên: ").bold = True
            fields.add_run(".................................................................")
            fields.add_run("     Số báo danh: ").bold = True
            fields.add_run(".....")
            fields.add_run(f"                 Mã đề : {exam_number:03d}").bold = True

            # Add bold horizontal line
            add_horizontal_line(fields)

            # doc.add_paragraph()  # Add some space

            # Add Part 1 questions to the document with custom heading
            heading = doc.add_paragraph(selected_headings["part1"], style=heading_style)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for i, (bai, phan, q_num, q_text) in enumerate(exam_questions_part1, 1):
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(0)
                p.add_run(f"Câu {i}: ").bold = True
                p.add_run(q_text.split('\n')[0])
                format_question(doc, '\n'.join(q_text.split('\n')[1:]), is_multiple_choice=True)

            doc.add_paragraph()  # Add some space
            # Add Part 2 questions to the document with custom heading
            heading = doc.add_paragraph(selected_headings["part2"], style=heading_style)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for i, (bai, phan, q_num, q_text) in enumerate(exam_questions_part2, 1):
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(0)
                p.add_run(f"Câu {i}: ").bold = True
                p.add_run(q_text.split('\n')[0])
                format_question(doc, '\n'.join(q_text.split('\n')[1:]), is_multiple_choice=False)

            # Count pages and update the header
            page_count = count_pages(doc)
            left_para.add_run(f"(Đề thi có {page_count} trang)").italic = True

            # Save Word file to buffer
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            # Add Word file to ZIP
            zip_file.writestr(f'de_thi_{exam_number:03d}.docx', docx_buffer.getvalue())

    # Prepare ZIP buffer for download
    zip_buffer.seek(0)

    # Create download link for ZIP file
    st.download_button(
        label="Tải xuống tất cả đề thi",
        data=zip_buffer,
        file_name="de_thi.zip",
        mime="application/zip"
    )